from GT_imports import *
import json
from tkinter.filedialog import asksaveasfilename, askopenfilename
from tkinter import messagebox, Text
from docx import Document  # For DOCX files

class FileOps:
    def __init__(self):
        pass
    
    # ----------------------- FILE OPERATIONS MANAGEMENT -----------------------------------------------------
    def _save_as(self):
        file_path = asksaveasfilename(
            defaultextension=".txt",
            filetypes=[
                ("Text Files", "*.txt"),
                ("JSON Files", "*.json"),
                ("RTF Files", "*.rtf"),
                ("Word Document", "*.docx"),
                ("All Files", "*.*")
            ],
            title="Save As"
        )
        
        if file_path:
            file_extension = file_path.split('.')[-1].lower()
            content = self._get_current_text_area().get("1.0", "end-1c")

            try:
                if file_extension == "txt":
                    with open(file_path, "w") as file:
                        file.write(content)

                elif file_extension == "json":
                    data = json.loads(content)  
                    with open(file_path, "w") as file:
                        json.dump(data, file, indent=4)

                elif file_extension == "rtf":
                    # Save as RTF (we'll wrap the text in basic RTF format)
                    with open(file_path, "w") as file:
                        rtf_content = "{\\rtf1\\ansi\\ansicpg1252\\deff0\\nouicompat\\deflang1033{\\fonttbl{\\f0\\fnil\\fcharset0 Calibri;}}"
                        rtf_content += "\\viewkind4\\uc1 \\pard\\sa200\\sl276\\slmult1\\f0\\fs22\\pardirnatural " + content.replace("\n", "\\par ") + "}"
                        file.write(rtf_content)

                elif file_extension == "docx":
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(file_path)

                else:
                    messagebox.showerror("Error", "Unsupported file type.")
            
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred while saving the file: {e}")

    def _load_text(self):
        """Open a file using a file dialog and load its content into the current tab."""
        file_path = askopenfilename(
            title="Open File",
            filetypes=[
                ("Text Files", "*.txt"),
                ("JSON Files", "*.json"),
                ("Word Documents", "*.docx"),
                ("All Files", "*.*")
            ]
        )
        
        if file_path:
            file_extension = file_path.split('.')[-1].lower()
            
            try:
                if file_extension == "txt":
                    with open(file_path, "r") as file:
                        content = file.read()
                    self._get_current_text_area().delete("1.0", "end")
                    self._get_current_text_area().insert("1.0", content)

                elif file_extension == "json":
                    with open(file_path, "r") as file:
                        content = json.load(file)
                    self._get_current_text_area().delete("1.0", "end")
                    self._get_current_text_area().insert("1.0", json.dumps(content, indent=4))

                elif file_extension == "docx":
                    doc = Document(file_path)
                    self._load_docx(doc)

                else:
                    messagebox.showerror("Error", "Unsupported file type.")
            
            except Exception as e:
                messagebox.showerror("Error", f"An error occurred while loading the file: {e}")

    def _load_docx(self, doc):
        """Load a DOCX file, maintaining formatting like bold, italic, and underline."""
        text_widget = self._get_current_text_area()
        text_widget.delete("1.0", "end")
        
        for para in doc.paragraphs:
            for run in para.runs:
                if run.bold:
                    text_widget.insert("end", f"<b>{run.text}</b>")
                elif run.italic:
                    text_widget.insert("end", f"<i>{run.text}</i>")
                elif run.underline:
                    text_widget.insert("end", f"<u>{run.text}</u>")
                else:
                    text_widget.insert("end", run.text)
            text_widget.insert("end", "\n") 

    def _show_about(self):
        messagebox.showinfo("About", "This is a basic text editor with tabs, find & replace, and more.") 

    # -----------------------END OF FILE OPERATIONS MANAGEMENT -----------------------------------------------------

